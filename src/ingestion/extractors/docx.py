"""
DOCX extractor using python-docx.
"""

import re
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_BREAK_TYPE
from docx.shared import Inches

from .base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """
    DOCX document extractor using python-docx.
    
    Extracts text and metadata from Word documents with configurable options
    for handling tables, images, headers/footers, and document structure.
    """
    
    def __init__(self):
        """Initialize the DOCX extractor."""
        self.document = None
    
    def can_handle(self, file_path: Path) -> bool:
        """
        Check if this extractor can handle DOCX files.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if file has .docx extension, False otherwise
        """
        return file_path.suffix.lower() == '.docx'
    
    def extract_text(self, file_path: Path, config: Dict[str, Any]) -> str:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            config: Configuration dictionary
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX cannot be opened or processed
        """
        try:
            self.document = Document(file_path)
            text_parts = []
            
            # Extract main document text
            if config.get('extract_paragraphs', True):
                text_parts.extend(self._extract_paragraphs(config))
            
            # Extract tables
            if config.get('extract_tables', True):
                table_text = self._extract_tables(config)
                if table_text:
                    text_parts.append(table_text)
            
            # Extract headers and footers
            if config.get('extract_headers_footers', True):
                header_footer_text = self._extract_headers_footers(config)
                if header_footer_text:
                    text_parts.append(header_footer_text)
            
            # Extract footnotes and endnotes
            if config.get('extract_notes', True):
                notes_text = self._extract_notes(config)
                if notes_text:
                    text_parts.append(notes_text)
            
            # Join all text parts
            text = '\n\n'.join(filter(None, text_parts))
            
            # Clean and normalize text
            text = self._clean_text(text, config)
            
            return text
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_paragraphs(self, config: Dict[str, Any]) -> List[str]:
        """
        Extract text from document paragraphs.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            List of paragraph texts
        """
        paragraphs = []
        preserve_formatting = config.get('preserve_formatting', True)
        include_styles = config.get('include_paragraph_styles', False)
        
        for paragraph in self.document.paragraphs:
            if not paragraph.text.strip():
                continue
            
            text = paragraph.text
            
            # Add style information if requested
            if include_styles and paragraph.style.name != 'Normal':
                text = f"[{paragraph.style.name}] {text}"
            
            # Handle formatting if requested
            if preserve_formatting:
                # Check for page breaks
                for run in paragraph.runs:
                    if run._element.xml.find('w:br') != -1:
                        # Check if it's a page break
                        if 'type="page"' in run._element.xml:
                            text += '\n\n[PAGE BREAK]\n\n'
                        elif 'type="column"' in run._element.xml:
                            text += '\n\n[COLUMN BREAK]\n\n'
            
            paragraphs.append(text)
        
        return paragraphs
    
    def _extract_tables(self, config: Dict[str, Any]) -> str:
        """
        Extract text from document tables.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            Formatted table text
        """
        if not self.document.tables:
            return ""
        
        table_texts = []
        table_format = config.get('table_format', 'grid')  # 'grid', 'simple', 'csv'
        
        for i, table in enumerate(self.document.tables):
            if config.get('include_table_headers', True):
                table_texts.append(f"\n[TABLE {i+1}]\n")
            
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    # Extract text from all paragraphs in the cell
                    cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    cells.append(cell_text)
                rows.append(cells)
            
            # Format table based on configuration
            if table_format == 'csv':
                formatted_table = self._format_table_csv(rows)
            elif table_format == 'simple':
                formatted_table = self._format_table_simple(rows)
            else:  # grid
                formatted_table = self._format_table_grid(rows)
            
            table_texts.append(formatted_table)
        
        return '\n\n'.join(table_texts)
    
    def _format_table_grid(self, rows: List[List[str]]) -> str:
        """Format table in grid format with borders."""
        if not rows:
            return ""
        
        # Calculate column widths
        col_widths = [0] * len(rows[0])
        for row in rows:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(cell))
        
        # Create grid format
        lines = []
        separator = '+' + '+'.join('-' * (w + 2) for w in col_widths) + '+'
        
        lines.append(separator)
        for row in rows:
            line = '|'
            for i, cell in enumerate(row):
                line += f' {cell:<{col_widths[i]}} |'
            lines.append(line)
            lines.append(separator)
        
        return '\n'.join(lines)
    
    def _format_table_simple(self, rows: List[List[str]]) -> str:
        """Format table in simple format with spaces."""
        if not rows:
            return ""
        
        # Calculate column widths
        col_widths = [0] * len(rows[0])
        for row in rows:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(cell))
        
        # Create simple format
        lines = []
        for row in rows:
            line = '  '.join(f'{cell:<{col_widths[i]}}' for i, cell in enumerate(row))
            lines.append(line)
        
        return '\n'.join(lines)
    
    def _format_table_csv(self, rows: List[List[str]]) -> str:
        """Format table in CSV format."""
        if not rows:
            return ""
        
        lines = []
        for row in rows:
            # Escape cells containing commas or quotes
            escaped_cells = []
            for cell in row:
                if ',' in cell or '"' in cell or '\n' in cell:
                    cell = '"' + cell.replace('"', '""') + '"'
                escaped_cells.append(cell)
            lines.append(','.join(escaped_cells))
        
        return '\n'.join(lines)
    
    def _extract_headers_footers(self, config: Dict[str, Any]) -> str:
        """
        Extract text from headers and footers.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            Header and footer text
        """
        header_footer_texts = []
        
        try:
            # Extract from all sections
            for section in self.document.sections:
                # Headers
                if section.header:
                    header_text = '\n'.join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                    if header_text:
                        header_footer_texts.append(f"[HEADER]\n{header_text}")
                
                # Footers
                if section.footer:
                    footer_text = '\n'.join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                    if footer_text:
                        header_footer_texts.append(f"[FOOTER]\n{footer_text}")
        except Exception as e:
            # Headers/footers might not be accessible in some documents
            import logging
            logger = logging.getLogger(__name__)
            logger.debug(f"Could not extract headers/footers: {e}")
        
        return '\n\n'.join(header_footer_texts)
    
    def _extract_notes(self, config: Dict[str, Any]) -> str:
        """
        Extract footnotes and endnotes.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            Notes text
        """
        # Note: python-docx has limited support for footnotes/endnotes
        # This is a placeholder for future enhancement
        return ""
    
    def _clean_text(self, text: str, config: Dict[str, Any]) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            config: Configuration dictionary
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        if config.get('normalize_whitespace', True):
            # Replace multiple spaces with single space
            text = re.sub(r' +', ' ', text)
            # Replace multiple newlines with double newline
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            # Remove leading/trailing whitespace from lines
            lines = [line.strip() for line in text.split('\n')]
            text = '\n'.join(lines)
        
        # Remove empty lines if configured
        if config.get('remove_empty_lines', True):
            lines = [line for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)
        
        return text.strip()
    
    def extract_metadata(self, file_path: Path, config: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            config: Configuration dictionary
            
        Returns:
            Dictionary containing DOCX metadata
            
        Raises:
            Exception: If DOCX cannot be opened or metadata cannot be extracted
        """
        try:
            # If document not already loaded, load it
            if self.document is None:
                self.document = Document(file_path)
            
            metadata = {
                'filename': file_path.name,
                'file_size': file_path.stat().st_size,
            }
            
            # Extract core properties
            core_props = self.document.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            if core_props.category:
                metadata['category'] = core_props.category
            if core_props.comments:
                metadata['comments'] = core_props.comments
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modification_date'] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            if core_props.revision:
                metadata['revision'] = core_props.revision
            
            # Document statistics
            if config.get('include_statistics', True):
                stats = self._extract_document_statistics()
                metadata.update(stats)
            
            # Document structure
            if config.get('include_structure', True):
                structure = self._extract_document_structure()
                metadata.update(structure)
            
            # Remove empty values
            metadata = {k: v for k, v in metadata.items() if v}
            
            return metadata
            
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            raise
    
    def _extract_document_statistics(self) -> Dict[str, Any]:
        """
        Extract document statistics.
        
        Returns:
            Dictionary with document statistics
        """
        stats = {}
        
        # Count paragraphs
        stats['paragraph_count'] = len([p for p in self.document.paragraphs if p.text.strip()])
        
        # Count tables
        stats['table_count'] = len(self.document.tables)
        
        # Count pages (approximate)
        # Note: Accurate page count requires more complex analysis
        char_count = sum(len(p.text) for p in self.document.paragraphs)
        stats['estimated_page_count'] = max(1, char_count // 2500)  # Rough estimate
        
        # Count words
        word_count = sum(len(p.text.split()) for p in self.document.paragraphs)
        stats['word_count'] = word_count
        
        # Count characters
        stats['character_count'] = char_count
        
        return stats
    
    def _extract_document_structure(self) -> Dict[str, Any]:
        """
        Extract document structure information.
        
        Returns:
            Dictionary with structure information
        """
        structure = {}
        
        # Extract headings
        headings = {}
        for paragraph in self.document.paragraphs:
            style_name = paragraph.style.name
            if style_name.startswith('Heading'):
                level = style_name.replace('Heading ', 'h')
                if level not in headings:
                    headings[level] = []
                if paragraph.text.strip():
                    headings[level].append(paragraph.text.strip())
        
        if headings:
            structure['headings'] = headings
        
        # Extract styles used
        styles_used = set()
        for paragraph in self.document.paragraphs:
            styles_used.add(paragraph.style.name)
        
        structure['styles_used'] = sorted(list(styles_used))
        
        # Section count
        structure['section_count'] = len(self.document.sections)
        
        return structure
    
    @property
    def supported_extensions(self) -> list[str]:
        """
        List of file extensions this extractor supports.
        
        Returns:
            List containing '.docx'
        """
        return ['.docx']
    
    @property
    def extractor_name(self) -> str:
        """
        Human-readable name of this extractor.
        
        Returns:
            Name of the DOCX extractor
        """
        return "DOCX Extractor (python-docx)"